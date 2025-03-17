import os
import logging
import sys
import json
from pathlib import Path
import time
import shutil
import traceback

# 导入pdf2docx模块 - 专业的PDF转Word工具
from pdf2docx import Converter

# 保留一些可能需要的辅助模块
import fitz  # PyMuPDF用于辅助分析
import docx
from docx import Document
from docx.shared import Pt, Mm
import cv2
import numpy as np
from PIL import Image

# 导入pdf2docx转换器
try:
    from utils.pdf2docx_converter import convert_pdf_to_word as pdf2docx_convert
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False
    logging.warning("pdf2docx_converter模块导入失败，将使用备选方法")

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFToWordConverter:
    """使用pdf2docx实现的PDF转Word转换器"""
    
    def __init__(self, pdf_path: Path, output_path: Path, session_dir: Path):
        """初始化转换器
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Word文件路径
            session_dir: 会话目录，用于存储临时文件和日志
        """
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.session_dir = session_dir
        
        # 创建必要的目录
        self.debug_dir = session_dir / "debug"
        self.images_dir = session_dir / "images"
        self.debug_dir.mkdir(exist_ok=True)
        self.images_dir.mkdir(exist_ok=True)
        
        # 设置调试日志
        self.debug_logger = self._setup_debug_logger()
        
        # 记录基本信息
        self.debug_logger.info(f"初始化PDF转Word转换器")
        self.debug_logger.info(f"PDF文件: {pdf_path}")
        self.debug_logger.info(f"输出Word文件: {output_path}")
        self.debug_logger.info(f"会话目录: {session_dir}")
        self.debug_logger.info(f"Python版本: {sys.version}")
        
        # 分析PDF基本信息
        self._analyze_pdf()
    
    def _setup_debug_logger(self):
        """设置调试日志记录器"""
        debug_logger = logging.getLogger(f"{__name__}.debug")
        debug_logger.setLevel(logging.DEBUG)
        
        # 添加文件处理器
        debug_log_file = self.debug_dir / "conversion_debug.log"
        file_handler = logging.FileHandler(debug_log_file, mode='w', encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        debug_logger.addHandler(file_handler)
        
        # 添加控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(formatter)
        debug_logger.addHandler(console_handler)
        
        return debug_logger
    
    def _analyze_pdf(self):
        """分析PDF文件基本信息"""
        try:
            # 使用PyMuPDF获取PDF基本信息
            pdf_doc = fitz.open(str(self.pdf_path))
            
            # 记录PDF基本信息
            self.page_count = len(pdf_doc)
            self.metadata = pdf_doc.metadata
            
            self.debug_logger.info(f"PDF页数: {self.page_count}")
            self.debug_logger.info(f"PDF元数据: {self.metadata}")
            
            # 导出PDF信息用于调试
            info = {
                "file_path": str(self.pdf_path),
                "page_count": self.page_count,
                "metadata": self.metadata,
                "pages": []
            }
            
            # 收集每页信息
            for page_num in range(self.page_count):
                page = pdf_doc[page_num]
                page_info = {
                    "page_number": page_num + 1,
                    "width": page.rect.width,
                    "height": page.rect.height,
                    "rotation": page.rotation,
                    "image_count": len(page.get_images()),
                }
                info["pages"].append(page_info)
            
            # 保存到JSON文件
            info_path = self.debug_dir / "pdf_info.json"
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(info, f, indent=2, ensure_ascii=False)
            
            # 关闭PDF文档
            pdf_doc.close()
            
        except Exception as e:
            self.debug_logger.error(f"分析PDF时出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
    
    def convert(self) -> Path:
        """执行PDF到Word的转换
        
        Returns:
            输出Word文件的路径
        """
        start_time = time.time()
        self.debug_logger.info(f"开始转换PDF到Word: {self.pdf_path} -> {self.output_path}")
        
        try:
            # 使用pdf2docx进行转换
            # 这个库专门设计用于保持布局一致性，包括表格和图片的精确还原
            cv = Converter(str(self.pdf_path))
            
            # 设置转换参数
            # 使用默认参数，这些参数已经针对布局保持进行了优化
            cv.convert(str(self.output_path), start=0, end=None)
            
            # 关闭转换器
            cv.close()
            
            # 验证转换结果
            self._verify_conversion()
            
            elapsed_time = time.time() - start_time
            self.debug_logger.info(f"转换完成，耗时: {elapsed_time:.2f}秒")
            logger.info(f"PDF转Word完成: {self.output_path}")
            
            return self.output_path
            
        except Exception as e:
            self.debug_logger.error(f"转换过程中出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
            
            # 尝试使用备用方法
            return self._fallback_conversion()
    
    def _verify_conversion(self):
        """验证转换结果"""
        try:
            # 检查输出文件是否存在
            if not os.path.exists(self.output_path):
                self.debug_logger.error("转换失败: 输出文件不存在")
                return False
            
            # 检查Word文档页数
            doc = Document(self.output_path)
            word_pages = len(doc.sections)
            
            self.debug_logger.info(f"Word文档页数: {word_pages}")
            self.debug_logger.info(f"原PDF页数: {self.page_count}")
            
            # 记录页数是否一致
            if word_pages != self.page_count:
                self.debug_logger.warning(f"页数不一致: PDF={self.page_count}, Word={word_pages}")
            else:
                self.debug_logger.info("页数一致 ✓")
            
            return True
        except Exception as e:
            self.debug_logger.error(f"验证转换结果时出错: {str(e)}")
            return False
    
    def _fallback_conversion(self) -> Path:
        """备用转换方法，当pdf2docx失败时使用"""
        self.debug_logger.info("使用备用方法进行转换")
        
        try:
            # 创建一个简单的Word文档，包含错误信息
            doc = Document()
            doc.add_heading("PDF转换失败", 0)
            doc.add_paragraph("使用pdf2docx转换PDF时出错。请尝试以下解决方案:")
            doc.add_paragraph("1. 确保PDF文件未加密且可以正常打开")
            doc.add_paragraph("2. 尝试使用其他PDF查看器打开并另存为PDF，然后重新转换")
            doc.add_paragraph("3. 如果PDF包含复杂元素，可能需要手动调整转换后的文档")
            
            # 保存文档
            fallback_path = self.output_path.with_name(f"{self.output_path.stem}_fallback{self.output_path.suffix}")
            doc.save(fallback_path)
            
            self.debug_logger.info(f"已创建备用文档: {fallback_path}")
            return fallback_path
            
        except Exception as e:
            self.debug_logger.error(f"备用转换也失败: {str(e)}")
            raise

# 主转换函数
async def advanced_convert_pdf_to_word(pdf_path, output_path, session_dir, style_config=None):
    """
    使用高级转换器将PDF转换为Word
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出Word文件路径
        session_dir: 会话目录
        style_config: 样式配置（可选，如果为None则自动检测）
        
    Returns:
        输出Word文件路径
    """
    try:
        # 优先使用pdf2docx模块进行转换
        if HAS_PDF2DOCX:
            logger.info("使用pdf2docx模块进行PDF到Word的转换")
            logger.info(f"PDF文件: {pdf_path}")
            logger.info(f"输出文件: {output_path}")
            
            try:
                # 使用pdf2docx_converter中的转换函数
                return await pdf2docx_convert(pdf_path, output_path, session_dir, style_config)
            except Exception as pdf2docx_error:
                logger.error(f"pdf2docx转换失败: {str(pdf2docx_error)}")
                logger.info("尝试使用备用方法")
        else:
            logger.info("pdf2docx模块不可用，使用备用方法")
            
        # 使用原始方法作为备用
        logger.info("使用备用转换方法")
        converter = PDFToWordConverter(pdf_path, output_path, session_dir)
        return converter.convert()
            
    except Exception as e:
        logger.error(f"高级转换过程中出错: {str(e)}")
        raise

# 为了向后兼容，提供函数别名
precise_convert_pdf_to_word = advanced_convert_pdf_to_word 