"""
Markdown到Word转换模块 - 使用Pandoc实现
确保格式一致、代码高亮、数学公式支持
"""

import os
import logging
import sys
import json
from pathlib import Path
import time
import traceback
import asyncio
import shutil
import subprocess
import tempfile
import requests
import urllib.parse
from typing import Optional, Tuple
import pypandoc

# 导入docx模块用于处理Word文档
import docx
from docx import Document
from docx.shared import Pt, Mm
from docx.shared import Inches

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 创建全局日志目录
LOGS_DIR = Path("logs")
LOGS_DIR.mkdir(exist_ok=True)

class MarkdownToWordConverter:
    """使用Pandoc实现的Markdown转Word转换器"""
    
    def __init__(self, md_path: Path, output_path: Path, session_dir: Path):
        """初始化转换器
        
        Args:
            md_path: Markdown文件路径
            output_path: 输出Word文件路径
            session_dir: 会话目录，用于存储临时文件
        """
        self.md_path = md_path
        self.output_path = output_path
        self.session_dir = session_dir
        
        # 创建必要的目录
        self.temp_dir = session_dir / "temp"
        self.temp_dir.mkdir(exist_ok=True)
        
        # 创建图片缓存目录
        self.image_cache_dir = self.temp_dir / "images"
        self.image_cache_dir.mkdir(exist_ok=True)
        
        # 设置调试日志
        self.debug_logger = self._setup_debug_logger()
        
        # 记录基本信息
        self.debug_logger.info(f"初始化Markdown转Word转换器")
        self.debug_logger.info(f"Markdown文件: {md_path}")
        self.debug_logger.info(f"输出Word文件: {output_path}")
        self.debug_logger.info(f"会话目录: {session_dir}")
        self.debug_logger.info(f"Python版本: {sys.version}")
        
        # 分析Markdown基本信息
        self._analyze_markdown()
    
    def _setup_debug_logger(self):
        """设置调试日志记录器"""
        debug_logger = logging.getLogger(f"{__name__}.debug")
        debug_logger.setLevel(logging.DEBUG)
        
        # 使用任务ID作为日志文件名的一部分，确保唯一性
        task_id = self.session_dir.name
        debug_log_file = LOGS_DIR / f"md_conversion_debug_{task_id}.log"
        
        # 添加文件处理器
        file_handler = logging.FileHandler(debug_log_file, mode='w', encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        debug_logger.addHandler(file_handler)
        
        # 添加控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(formatter)
        debug_logger.addHandler(console_handler)
        
        return debug_logger
    
    def _analyze_markdown(self):
        """分析Markdown文件基本信息"""
        try:
            # 读取Markdown文件内容
            with open(self.md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 计算行数和字符数
            lines = content.split('\n')
            line_count = len(lines)
            char_count = len(content)
            
            # 检测是否包含代码块和数学公式
            has_code_blocks = '```' in content
            has_math = '$' in content
            
            # 记录Markdown基本信息
            self.debug_logger.info(f"Markdown行数: {line_count}")
            self.debug_logger.info(f"Markdown字符数: {char_count}")
            self.debug_logger.info(f"包含代码块: {has_code_blocks}")
            self.debug_logger.info(f"包含数学公式: {has_math}")
            
            # 导出Markdown信息用于调试
            info = {
                "file_path": str(self.md_path),
                "line_count": line_count,
                "char_count": char_count,
                "has_code_blocks": has_code_blocks,
                "has_math": has_math
            }
            
            # 保存到JSON文件
            info_path = self.temp_dir / "markdown_info.json"
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(info, f, indent=2, ensure_ascii=False)
            
        except Exception as e:
            self.debug_logger.error(f"分析Markdown时出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
    
    def _check_pandoc_installed(self) -> bool:
        """检查Pandoc是否已安装"""
        try:
            # 首先尝试使用pypandoc获取pandoc版本
            try:
                version = pypandoc.get_pandoc_version()
                self.debug_logger.info(f"Pandoc版本(pypandoc): {version}")
                return True
            except OSError:
                self.debug_logger.warning("通过pypandoc无法获取Pandoc版本，尝试直接检查pandoc命令")
            
            # 如果pypandoc方法失败，尝试直接运行pandoc命令
            pandoc_paths = [
                'pandoc',  # 系统PATH中的pandoc
                str(Path(sys.executable).parent / 'Scripts' / 'pandoc.exe'),  # Python环境中的pandoc
                str(Path.home() / '.local' / 'bin' / 'pandoc'),  # 用户本地安装的pandoc
            ]
            
            for pandoc_path in pandoc_paths:
                try:
                    result = subprocess.run([pandoc_path, '--version'], 
                                        stdout=subprocess.PIPE, 
                                        stderr=subprocess.PIPE, 
                                        text=True, 
                                        check=False)
                    if result.returncode == 0:
                        version = result.stdout.split('\n')[0]
                        self.debug_logger.info(f"Pandoc版本(命令行): {version}")
                        return True
                except Exception:
                    continue
            
            # 如果以上方法都失败，尝试下载并安装pandoc
            self.debug_logger.warning("Pandoc未找到，尝试自动下载安装")
            pypandoc.download_pandoc()
            
            # 再次验证安装
            version = pypandoc.get_pandoc_version()
            self.debug_logger.info(f"已安装Pandoc版本: {version}")
            return True
            
        except Exception as e:
            self.debug_logger.error(f"检查Pandoc安装时出错: {str(e)}")
            return False
    
    def _process_markdown_content(self, content: str) -> str:
        """处理Markdown内容，包括下载和替换图片路径
        
        Args:
            content: 原始Markdown内容
            
        Returns:
            处理后的Markdown内容
        """
        import re
        
        # 匹配Markdown图片语法
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        
        # 同步处理图片
        def process_image_sync(image_url: str):
            try:
                # 解析URL
                parsed_url = urllib.parse.urlparse(image_url)
                
                # 如果是网络图片
                if parsed_url.scheme in ('http', 'https'):
                    # 生成缓存文件名
                    cache_filename = Path(parsed_url.path).name
                    if not cache_filename:
                        cache_filename = f"image_{hash(image_url)}.png"
                    cache_path = self.image_cache_dir / cache_filename
                    
                    # 如果缓存中已存在，直接返回
                    if cache_path.exists():
                        return cache_path, None
                    
                    # 下载图片
                    try:
                        response = requests.get(image_url, timeout=10)
                        response.raise_for_status()
                        with open(cache_path, 'wb') as f:
                            f.write(response.content)
                        return cache_path, None
                    except Exception as e:
                        return None, f"下载图片失败: {str(e)}"
                
                # 如果是本地图片
                else:
                    local_path = Path(image_url)
                    # 如果是相对路径，相对于Markdown文件所在目录
                    if not local_path.is_absolute():
                        local_path = self.md_path.parent / local_path
                    
                    if local_path.exists():
                        # 复制到缓存目录
                        cache_path = self.image_cache_dir / local_path.name
                        shutil.copy2(local_path, cache_path)
                        return cache_path, None
                    else:
                        return None, f"本地图片不存在: {local_path}"
                    
            except Exception as e:
                return None, f"处理图片时出错: {str(e)}"
        
        # 同步处理所有图片
        result = content
        for match in re.finditer(image_pattern, content):
            alt_text, image_url = match.groups()
            image_path, error = process_image_sync(image_url.strip())
            
            if image_path:
                # 使用相对路径替换原始URL
                rel_path = os.path.relpath(image_path, self.md_path.parent)
                result = result.replace(match.group(0), f'![{alt_text}]({rel_path})')
            else:
                self.debug_logger.warning(f"处理图片失败: {error}")
                # 保留原始图片标记，但添加警告注释
                result = result.replace(match.group(0), 
                                     f'{match.group(0)}\n<!-- 警告：图片处理失败 - {error} -->')
        
        return result

    def convert(self) -> Path:
        """执行Markdown到Word的转换
        
        Returns:
            输出Word文件的路径
        """
        start_time = time.time()
        self.debug_logger.info(f"开始转换Markdown到Word: {self.md_path} -> {self.output_path}")
        
        try:
            # 检查Pandoc是否已安装
            if not self._check_pandoc_installed():
                self.debug_logger.error("Pandoc未安装，无法进行转换")
                return self._fallback_conversion("Pandoc未安装，请先安装Pandoc")
            
            # 读取并处理Markdown内容
            with open(self.md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 处理图片
            processed_content = self._process_markdown_content(content)
            
            # 保存处理后的Markdown内容到临时文件
            temp_md = self.temp_dir / "processed.md"
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            
            # 创建参考样式文件
            reference_docx = self._create_reference_docx(self.temp_dir)
            
            # 构建Pandoc命令
            cmd = [
                'pandoc',
                str(temp_md),
                '-f', 'markdown',
                '-t', 'docx',
                '-o', str(self.output_path),
                '--reference-doc', str(reference_docx) if reference_docx else '',
                '--standalone',
                '--highlight-style', 'tango',  # 代码高亮样式
                '--toc',  # 生成目录
                '--toc-depth', '3',  # 目录深度
                '--wrap', 'auto',  # 自动换行
                '--mathml',  # 使用MathML处理数学公式
                '--extract-media', str(self.image_cache_dir)  # 提取图片到指定目录
            ]
            
            # 移除空参数
            cmd = [arg for arg in cmd if arg]
            
            # 执行Pandoc命令
            self.debug_logger.info(f"执行命令: {' '.join(cmd)}")
            result = subprocess.run(cmd, 
                                stdout=subprocess.PIPE, 
                                stderr=subprocess.PIPE, 
                                text=True, 
                                check=False)
            
            if result.returncode == 0:
                self.debug_logger.info("Pandoc转换成功")
                
                # 验证转换结果
                if self._verify_conversion():
                    elapsed_time = time.time() - start_time
                    self.debug_logger.info(f"转换完成，耗时: {elapsed_time:.2f}秒")
                    logger.info(f"Markdown转Word完成: {self.output_path}")
                    return self.output_path
                else:
                    self.debug_logger.error("转换结果验证失败")
                    return self._fallback_conversion("转换结果验证失败")
            else:
                self.debug_logger.error(f"Pandoc转换失败: {result.stderr}")
                return self._fallback_conversion(f"Pandoc转换失败: {result.stderr}")
            
        except Exception as e:
            self.debug_logger.error(f"转换过程中出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
            
            # 尝试使用备用方法
            return self._fallback_conversion(str(e))
    
    def _create_reference_docx(self, temp_dir: Path) -> Path:
        """创建参考样式Word文档
        
        Args:
            temp_dir: 临时目录
            
        Returns:
            参考样式文档路径
        """
        try:
            # 创建一个包含基本样式的Word文档
            doc = Document()
            
            # 设置文档样式
            styles = doc.styles
            
            # 设置正文样式
            style = styles['Normal']
            font = style.font
            font.name = '微软雅黑'
            font.size = Pt(11)
            
            # 设置标题样式
            for i in range(1, 4):
                style_name = f'Heading {i}'
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = '微软雅黑'
                    font.size = Pt(16 - (i-1)*2)  # 一级标题16pt，二级14pt，三级12pt
                    font.bold = True
            
            # 设置代码块样式
            if 'Code' in styles:
                style = styles['Code']
                font = style.font
                font.name = 'Consolas'
                font.size = Pt(10)
            
            # 保存参考文档
            reference_path = temp_dir / "reference.docx"
            doc.save(reference_path)
            
            self.debug_logger.info(f"创建参考样式文档: {reference_path}")
            return reference_path
            
        except Exception as e:
            self.debug_logger.error(f"创建参考样式文档时出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
            
            # 如果创建失败，返回None，后续会使用Pandoc默认样式
            return None
    
    def _verify_conversion(self) -> bool:
        """验证转换结果"""
        try:
            # 检查输出文件是否存在
            if not self.output_path.exists():
                self.debug_logger.error("转换失败: 输出文件不存在")
                return False
            
            # 检查输出文件大小
            file_size = self.output_path.stat().st_size
            if file_size == 0:
                self.debug_logger.error("转换失败: 输出文件为空")
                return False
            
            # 尝试打开Word文档
            try:
                doc = Document(self.output_path)
                paragraph_count = len(doc.paragraphs)
                self.debug_logger.info(f"Word文档段落数: {paragraph_count}")
                
                # 如果段落数为0，可能转换有问题
                if paragraph_count == 0:
                    self.debug_logger.warning("Word文档不包含任何段落")
                    return False
                
                return True
            except Exception as e:
                self.debug_logger.error(f"验证Word文档时出错: {str(e)}")
                return False
            
        except Exception as e:
            self.debug_logger.error(f"验证转换结果时出错: {str(e)}")
            return False
    
    def _fallback_conversion(self, error_message: str) -> Path:
        """备用转换方法，当Pandoc失败时使用
        
        Args:
            error_message: 错误信息
            
        Returns:
            输出Word文件的路径
        """
        self.debug_logger.info("使用备用方法进行转换")
        
        try:
            # 创建一个简单的Word文档
            doc = Document()
            doc.add_heading("Markdown转换结果", 0)
            
            # 如果是Pandoc错误，添加错误信息
            if "Pandoc" in error_message:
                doc.add_paragraph(f"注意: 使用Pandoc转换时遇到问题: {error_message}")
                doc.add_paragraph("已使用备用方法转换，但格式可能不完整。")
            
            # 添加分隔线
            doc.add_paragraph("---")
            
            # 读取并处理Markdown内容
            try:
                with open(self.md_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # 处理图片
                processed_content = self._process_markdown_content(content)
                
                # 添加标题
                import re
                # 查找所有标题
                heading_pattern = r'^(#{1,6})\s+(.+)$'
                for line in processed_content.split('\n'):
                    match = re.match(heading_pattern, line)
                    if match:
                        level = len(match.group(1))
                        text = match.group(2)
                        doc.add_heading(text, level)
                
                # 添加正文内容（简单处理）
                paragraphs = processed_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.startswith('#'):
                        doc.add_paragraph(para)
                
                # 尝试添加图片
                image_pattern = r'!\[(.*?)\]\((.*?)\)'
                for match in re.finditer(image_pattern, processed_content):
                    alt_text, image_path = match.groups()
                    try:
                        # 检查图片路径是否在缓存目录中
                        img_path = Path(image_path)
                        if not img_path.is_absolute():
                            img_path = self.image_cache_dir / img_path.name
                        
                        if img_path.exists():
                            doc.add_picture(str(img_path), width=Inches(6))
                            doc.add_paragraph(f"图片说明: {alt_text}")
                    except Exception as img_e:
                        self.debug_logger.warning(f"添加图片失败: {str(img_e)}")
                        doc.add_paragraph(f"[无法添加图片: {alt_text}]")
                
            except Exception as e:
                self.debug_logger.error(f"无法读取原始Markdown内容: {str(e)}")
                doc.add_paragraph(f"无法读取原始Markdown内容: {str(e)}")
            
            # 保存文档
            fallback_path = self.output_path.with_name(f"{self.output_path.stem}_fallback{self.output_path.suffix}")
            doc.save(fallback_path)
            
            # 尝试复制到原始输出路径
            try:
                shutil.copy2(fallback_path, self.output_path)
                self.debug_logger.info(f"已将备用文档复制到原始输出路径: {self.output_path}")
                self.debug_logger.info(f"已创建备用文档: {fallback_path}")
                return self.output_path
            except Exception as copy_e:
                self.debug_logger.error(f"复制备用文档时出错: {str(copy_e)}")
                return fallback_path
            
        except Exception as e:
            self.debug_logger.error(f"备用转换也失败: {str(e)}")
            raise

# 主转换函数
async def convert_markdown_to_word(md_path, output_path, session_dir, style_config=None):
    """
    使用Pandoc将Markdown转换为Word，支持代码高亮和数学公式
    
    Args:
        md_path: Markdown文件路径
        output_path: 输出Word文件路径
        session_dir: 会话目录
        style_config: 样式配置（可选，目前未使用）
        
    Returns:
        输出Word文件路径
    """
    try:
        logger.info(f"开始转换Markdown到Word: {md_path} -> {output_path}")
        
        # 确保所有路径都是Path对象
        md_path = Path(md_path) if not isinstance(md_path, Path) else md_path
        output_path = Path(output_path) if not isinstance(output_path, Path) else output_path
        session_dir = Path(session_dir) if not isinstance(session_dir, Path) else session_dir
        
        # 创建logs目录（如果不存在）
        logs_dir = Path("logs")
        logs_dir.mkdir(exist_ok=True)
        
        # 创建转换器并执行转换
        converter = MarkdownToWordConverter(md_path, output_path, session_dir)
        result_path = converter.convert()
        
        # 如果转换成功但生成的是备用文件，复制到原始输出路径
        if result_path != output_path and result_path.exists():
            try:
                shutil.copy2(result_path, output_path)
                logger.info(f"已将备用文件复制到原始输出路径: {output_path}")
                # 确保原始输出路径存在
                if output_path.exists():
                    result_path = output_path
            except Exception as copy_e:
                logger.error(f"复制备用文件时出错: {str(copy_e)}")
        
        # 设置定时任务，延迟5秒后删除临时文件，给前端足够时间下载
        asyncio.create_task(_delayed_cleanup(str(session_dir), 5))
        
        logger.info(f"转换完成: {result_path}")
        return result_path
        
    except Exception as e:
        logger.error(f"转换过程中出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise

async def _delayed_cleanup(session_dir: str, delay_seconds: int = 3):
    """
    延迟指定秒数后删除会话目录
    
    Args:
        session_dir: 会话目录路径
        delay_seconds: 延迟秒数
    """
    try:
        logger.info(f"计划在 {delay_seconds} 秒后清理临时文件: {session_dir}")
        await asyncio.sleep(delay_seconds)
        
        # 删除会话目录及其所有内容
        if os.path.exists(session_dir):
            # 关闭所有日志处理器
            task_id = os.path.basename(session_dir)
            debug_logger_name = f"{__name__}.debug"
            debug_logger = logging.getLogger(debug_logger_name)
            
            # 关闭并移除所有处理器
            handlers = debug_logger.handlers.copy()
            for handler in handlers:
                try:
                    handler.close()
                    debug_logger.removeHandler(handler)
                    logger.info(f"已关闭日志处理器: {handler}")
                except Exception as h_e:
                    logger.warning(f"关闭日志处理器时出错: {str(h_e)}")
            
            # 直接删除会话目录
            try:
                shutil.rmtree(session_dir, ignore_errors=True)
                logger.info(f"已删除会话目录: {session_dir}")
            except Exception as e:
                logger.warning(f"删除会话目录时出错: {str(e)}")
                # 如果删除失败，尝试逐个删除文件
                try:
                    for root, dirs, files in os.walk(session_dir, topdown=False):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                # 尝试修改文件权限
                                os.chmod(file_path, 0o777)
                                os.remove(file_path)
                                logger.info(f"已删除文件: {file_path}")
                            except Exception as file_e:
                                logger.warning(f"删除文件时出错: {file_path}, {str(file_e)}")
                        
                        for dir in dirs:
                            dir_path = os.path.join(root, dir)
                            try:
                                os.rmdir(dir_path)
                                logger.info(f"已删除目录: {dir_path}")
                            except Exception as dir_e:
                                logger.warning(f"删除目录时出错: {dir_path}, {str(dir_e)}")
                    
                    # 最后尝试删除根目录
                    try:
                        os.rmdir(session_dir)
                        logger.info(f"已删除根目录: {session_dir}")
                    except Exception as root_e:
                        logger.warning(f"删除根目录时出错: {session_dir}, {str(root_e)}")
                except Exception as walk_e:
                    logger.error(f"遍历删除文件时出错: {str(walk_e)}")
    except Exception as e:
        logger.error(f"清理临时文件时出错: {str(e)}") 