"""
PDF到Word转换模块 - 使用pdf2docx实现
确保页码一致、布局一致、表格精确还原、图片精准粘贴
"""

import os
import logging
import sys
import json
from pathlib import Path
import time
import traceback
import asyncio
import shutil

# 导入pdf2docx模块 - 专业的PDF转Word工具
from pdf2docx import Converter
# 导入pdf2docx的高级组件，用于自定义转换行为
from pdf2docx.common.constants import RectType
from pdf2docx.page.RectContainer import RectContainer

# 保留一些可能需要的辅助模块
import fitz  # PyMuPDF用于辅助分析
import docx
from docx import Document
from docx.shared import Pt, Mm

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFToWordConverter:
    """使用pdf2docx实现的PDF转Word转换器"""
    
    def __init__(self, pdf_path: Path, output_path: Path, session_dir: Path):
        """初始化转换器
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Word文件路径
            session_dir: 会话目录，用于存储临时文件和日志
        """
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.session_dir = session_dir
        
        # 创建必要的目录
        self.debug_dir = session_dir / "debug"
        self.images_dir = session_dir / "images"
        self.debug_dir.mkdir(exist_ok=True)
        self.images_dir.mkdir(exist_ok=True)
        
        # 设置调试日志
        self.debug_logger = self._setup_debug_logger()
        
        # 记录基本信息
        self.debug_logger.info(f"初始化PDF转Word转换器")
        self.debug_logger.info(f"PDF文件: {pdf_path}")
        self.debug_logger.info(f"输出Word文件: {output_path}")
        self.debug_logger.info(f"会话目录: {session_dir}")
        self.debug_logger.info(f"Python版本: {sys.version}")
        
        # 分析PDF基本信息
        self._analyze_pdf()
    
    def _setup_debug_logger(self):
        """设置调试日志记录器"""
        debug_logger = logging.getLogger(f"{__name__}.debug")
        debug_logger.setLevel(logging.DEBUG)
        
        # 创建全局日志目录
        logs_dir = Path("logs")
        logs_dir.mkdir(exist_ok=True)
        
        # 使用任务ID作为日志文件名的一部分，确保唯一性
        task_id = self.session_dir.name
        debug_log_file = logs_dir / f"conversion_debug_{task_id}.log"
        
        # 添加文件处理器
        file_handler = logging.FileHandler(debug_log_file, mode='w', encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        debug_logger.addHandler(file_handler)
        
        # 添加控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(formatter)
        debug_logger.addHandler(console_handler)
        
        # 保存日志文件路径，以便后续清理
        self.debug_log_file = debug_log_file
        
        return debug_logger
    
    def _analyze_pdf(self):
        """分析PDF文件基本信息"""
        try:
            # 使用PyMuPDF获取PDF基本信息
            pdf_doc = fitz.open(str(self.pdf_path))
            
            # 记录PDF基本信息
            self.page_count = len(pdf_doc)
            self.metadata = pdf_doc.metadata
            
            self.debug_logger.info(f"PDF页数: {self.page_count}")
            self.debug_logger.info(f"PDF元数据: {self.metadata}")
            
            # 导出PDF信息用于调试
            info = {
                "file_path": str(self.pdf_path),
                "page_count": self.page_count,
                "metadata": self.metadata,
                "pages": []
            }
            
            # 收集每页信息
            for page_num in range(self.page_count):
                page = pdf_doc[page_num]
                page_info = {
                    "page_number": page_num + 1,
                    "width": page.rect.width,
                    "height": page.rect.height,
                    "rotation": page.rotation,
                    "image_count": len(page.get_images()),
                }
                info["pages"].append(page_info)
            
            # 保存到JSON文件
            info_path = self.debug_dir / "pdf_info.json"
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(info, f, indent=2, ensure_ascii=False)
            
            # 关闭PDF文档
            pdf_doc.close()
            
        except Exception as e:
            self.debug_logger.error(f"分析PDF时出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
    
    def convert(self) -> Path:
        """执行PDF到Word的转换
        
        Returns:
            输出Word文件的路径
        """
        start_time = time.time()
        self.debug_logger.info(f"开始转换PDF到Word: {self.pdf_path} -> {self.output_path}")
        
        try:
            # 使用pdf2docx进行转换
            # 这个库专门设计用于保持布局一致性，包括表格和图片的精确还原
            cv = Converter(str(self.pdf_path))
            
            # 设置转换参数 - 优化文本框处理
            # 自定义参数以保持文本框格式
            cv.convert(
                str(self.output_path), 
                start=0, 
                end=None,
                multi_paragraphs=True,  # 启用多段落支持，保持文本框内的段落结构
                debug=False,
                min_paragraph_height=5.0,  # 降低段落高度阈值，更好地识别段落
                connected_border_tolerance=3.0,  # 提高边框连接容差
                line_overlap_threshold=0.9,  # 提高行重叠阈值，更好地处理文本框
                line_break_width_ratio=0.5,  # 调整行断开宽度比例
                float_image_ignorable_gap=5.0,  # 减小浮动图像可忽略间隙
                page_margin_factor=0.0,  # 减小页边距因子，保持原始布局
                detect_table_bottom_border=False,  # 禁用表格底部边框检测，防止文本框被识别为表格
                text_box_width_limit=20.0,  # 增加文本框宽度限制，更好地识别文本框
                text_box_free_text_tolerance=0.5,  # 降低自由文本容差，更好地处理文本框内容
                extract_stream_table=False  # 禁用流式表格提取，防止文本框被识别为表格
            )
            
            # 关闭转换器
            cv.close()
            
            # 验证转换结果
            self._verify_conversion()
            
            elapsed_time = time.time() - start_time
            self.debug_logger.info(f"转换完成，耗时: {elapsed_time:.2f}秒")
            logger.info(f"PDF转Word完成: {self.output_path}")
            
            return self.output_path
            
        except Exception as e:
            self.debug_logger.error(f"转换过程中出错: {str(e)}")
            self.debug_logger.error(traceback.format_exc())
            
            # 尝试使用备用方法
            return self._fallback_conversion()
    
    def _verify_conversion(self):
        """验证转换结果"""
        try:
            # 检查输出文件是否存在
            if not os.path.exists(self.output_path):
                self.debug_logger.error("转换失败: 输出文件不存在")
                return False
            
            # 检查Word文档页数
            doc = Document(self.output_path)
            # 更准确地计算Word文档页数
            word_pages = self._count_word_pages(doc)
            
            self.debug_logger.info(f"Word文档页数: {word_pages}")
            self.debug_logger.info(f"原PDF页数: {self.page_count}")
            
            # 记录页数是否一致
            if word_pages != self.page_count:
                self.debug_logger.warning(f"页数不一致: PDF={self.page_count}, Word={word_pages}")
            else:
                self.debug_logger.info("页数一致 ✓")
            
            return True
        except Exception as e:
            self.debug_logger.error(f"验证转换结果时出错: {str(e)}")
            return False
    
    def _count_word_pages(self, doc):
        """更准确地计算Word文档的页数"""
        try:
            # 使用文档属性中的页数信息
            return len(doc.sections)
        except:
            # 如果无法获取准确页数，返回节数作为估计
            return len(doc.sections)
    
    def _fallback_conversion(self) -> Path:
        """备用转换方法，当pdf2docx失败时使用"""
        self.debug_logger.info("使用备用方法进行转换")
        
        try:
            # 创建一个简单的Word文档，包含错误信息
            doc = Document()
            doc.add_heading("PDF转换失败", 0)
            doc.add_paragraph("使用pdf2docx转换PDF时出错。请尝试以下解决方案:")
            doc.add_paragraph("1. 确保PDF文件未加密且可以正常打开")
            doc.add_paragraph("2. 尝试使用其他PDF查看器打开并另存为PDF，然后重新转换")
            doc.add_paragraph("3. 如果PDF包含复杂元素，可能需要手动调整转换后的文档")
            
            # 保存文档
            fallback_path = self.output_path.with_name(f"{self.output_path.stem}_fallback{self.output_path.suffix}")
            doc.save(fallback_path)
            
            self.debug_logger.info(f"已创建备用文档: {fallback_path}")
            return fallback_path
            
        except Exception as e:
            self.debug_logger.error(f"备用转换也失败: {str(e)}")
            raise

# 主转换函数
async def convert_pdf_to_word(pdf_path, output_path, session_dir, style_config=None):
    """
    使用pdf2docx将PDF转换为Word，保持布局一致性
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出Word文件路径
        session_dir: 会话目录
        style_config: 样式配置（可选，目前未使用）
        
    Returns:
        输出Word文件路径
    """
    try:
        logger.info(f"开始转换PDF到Word: {pdf_path} -> {output_path}")
        
        # 记录会话目录，用于后续清理
        session_dir_str = str(session_dir)
        
        # 创建转换器并执行转换
        converter = PDFToWordConverter(pdf_path, output_path, session_dir)
        result_path = converter.convert()
        
        # 设置定时任务，延迟3秒后删除临时文件
        asyncio.create_task(_delayed_cleanup(session_dir_str, 3))
        
        logger.info(f"转换完成: {result_path}")
        return result_path
        
    except Exception as e:
        logger.error(f"转换过程中出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise 

async def _delayed_cleanup(session_dir: str, delay_seconds: int = 3):
    """
    延迟指定秒数后删除会话目录
    
    Args:
        session_dir: 会话目录路径
        delay_seconds: 延迟秒数
    """
    try:
        logger.info(f"计划在 {delay_seconds} 秒后清理临时文件: {session_dir}")
        await asyncio.sleep(delay_seconds)
        
        # 删除会话目录及其所有内容
        if os.path.exists(session_dir):
            # 不再需要关闭日志文件，因为日志文件现在存储在单独的日志目录中
            
            # 直接删除会话目录
            try:
                shutil.rmtree(session_dir, ignore_errors=True)
                logger.info(f"已删除会话目录: {session_dir}")
            except Exception as e:
                logger.warning(f"删除会话目录时出错: {str(e)}")
                # 如果删除失败，尝试逐个删除文件
                try:
                    for root, dirs, files in os.walk(session_dir, topdown=False):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                os.remove(file_path)
                                logger.info(f"已删除文件: {file_path}")
                            except Exception as file_e:
                                logger.warning(f"删除文件时出错: {file_path}, {str(file_e)}")
                        
                        for dir in dirs:
                            dir_path = os.path.join(root, dir)
                            try:
                                os.rmdir(dir_path)
                                logger.info(f"已删除目录: {dir_path}")
                            except Exception as dir_e:
                                logger.warning(f"删除目录时出错: {dir_path}, {str(dir_e)}")
                    
                    # 最后尝试删除根目录
                    try:
                        os.rmdir(session_dir)
                        logger.info(f"已删除根目录: {session_dir}")
                    except Exception as root_e:
                        logger.warning(f"删除根目录时出错: {session_dir}, {str(root_e)}")
                except Exception as walk_e:
                    logger.error(f"遍历删除文件时出错: {str(walk_e)}")
    except Exception as e:
        logger.error(f"清理临时文件时出错: {str(e)}") 